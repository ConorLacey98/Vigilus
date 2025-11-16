# core/reporting.py
from docx import Document
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any

from .db import get_connection  # reuse your DB helper

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "reports" / "company_advisory_template.docx"
OUTPUT_DIR = Path(__file__).resolve().parent.parent / "generated_reports"
OUTPUT_DIR.mkdir(exist_ok=True)


def _replace_placeholder(doc: Document, placeholder: str, value: str) -> None:
    """Simple global find/replace for placeholders in all paragraphs and table cells."""
    if value is None:
        value = ""

    for p in doc.paragraphs:
        if placeholder in p.text:
            for run in p.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)


def _get_event_with_sources(event_id: int) -> Dict[str, Any]:
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("SELECT * FROM events WHERE id = ?", (event_id,))
    event = cur.fetchone()
    if not event:
        conn.close()
        raise ValueError(f"No event found with id {event_id}")

    cur.execute(
        """
        SELECT ri.*
        FROM event_sources es
        JOIN raw_items ri ON es.raw_item_id = ri.id
        WHERE es.event_id = ?
        """,
        (event_id,),
    )
    sources = cur.fetchall()
    conn.close()
    return {"event": event, "sources": sources}


def generate_advisory_docx(event_id: int) -> Path:
    data = _get_event_with_sources(event_id)
    ev = data["event"]
    sources = data["sources"]

    doc = Document(TEMPLATE_PATH)

    # Build fields
    title = ev["summary"] or "Security Advisory"
    date_str = datetime.utcnow().strftime("%Y-%m-%d")
    severity = ev["severity"] or "TBD"
    vendor = ev["vendor"] or ""
    cves = (ev["cves"] or "").strip()
    summary = ev["summary"] or ""

    details_lines = []
    for s in sources:
        line = f"- [{s['source']}] {s['title'] or ''}".strip()
        if s["url"]:
            line += f" ({s['url']})"
        details_lines.append(line)
    details_text = "\n".join(details_lines) if details_lines else "Further technical details pending."

    recommendations_text = (
        "- Identify affected systems.\n"
        "- Apply vendor patches or mitigations.\n"
        "- Review logs for signs of exploitation.\n"
        "- Update detection rules as needed."
    )

    references_lines = [s["url"] for s in sources if s["url"]]
    references_text = "\n".join(references_lines)

    # Replace placeholders
    _replace_placeholder(doc, "{{TITLE}}", title)
    _replace_placeholder(doc, "{{DATE}}", date_str)
    _replace_placeholder(doc, "{{SEVERITY}}", severity)
    _replace_placeholder(doc, "{{VENDOR}}", vendor)
    _replace_placeholder(doc, "{{CVES}}", cves)
    _replace_placeholder(doc, "{{SUMMARY}}", summary)
    _replace_placeholder(doc, "{{DETAILS}}", details_text)
    _replace_placeholder(doc, "{{RECOMMENDATIONS}}", recommendations_text)
    _replace_placeholder(doc, "{{REFERENCES}}", references_text)

    # Save
    safe_title = "".join(c for c in title if c.isalnum() or c in (" ", "_", "-"))[:80].strip().replace(" ", "_")
    filename = f"Advisory_{safe_title}_{date_str}.docx"
    output_path = OUTPUT_DIR / filename
    doc.save(output_path)

    return output_path